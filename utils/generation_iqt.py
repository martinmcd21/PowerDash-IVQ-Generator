# utils/export_iqt.py
import io, os, requests
from typing import Dict

# ---------- DOCX ----------
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- PDF ----------
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader

FONT_NAME = "Source Sans 3"

# ==============================
# DOCX helpers
# ==============================
def _set_document_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(11)

def _ruled_paragraph(doc, width_chars: int = 100):
    # A dashed underline line for handwritten notes
    p = doc.add_paragraph(" ")
    p.paragraph_format.space_after = Pt(2)
    p_pr = p._element.get_or_add_pPr()
    p_brd = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dashed")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")
    p_brd.append(bottom)
    p_pr.append(p_brd)

def _add_footer_powerdash(doc: Document, pd_logo_path: str):
    # Footer on every section → repeats on every page
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        run = p.add_run()
        try:
            if pd_logo_path and os.path.exists(pd_logo_path):
                run.add_picture(pd_logo_path, width=Inches(0.6))
                p.add_run("  Powered by PowerDash HR").italic = True
            else:
                p.add_run("Powered by PowerDash HR").italic = True
        except Exception:
            p.add_run("Powered by PowerDash HR").italic = True

def _add_question_table(doc: Document, q: Dict):
    """
    DOCX: Question uses a full-width row (merged across both columns),
    then label/value rows for Intent / Follow-ups / What good looks like,
    followed by 5 ruled note lines.
    """
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False

    # Set explicit column widths (label narrow, value wide)
    try:
        tbl.columns[0].width = Inches(1.2)
        tbl.columns[1].width = Inches(5.8)
    except Exception:
        pass

    # --- QUESTION (full width) ---
    row = tbl.add_row().cells
    try:
        qcell = row[0].merge(row[1])
    except Exception:
        qcell = row[0]
    q_para = qcell.paragraphs[0]
    q_run = q_para.add_run((q.get("question") or "").strip())
    q_run.bold = True
    q_para.space_after = Pt(6)

    # helper for label/value rows
    def add_row(label, val):
        if not val:
            return
        r = tbl.add_row().cells
        lbl_para = r[0].paragraphs[0]; lbl_run = lbl_para.add_run(label); lbl_run.bold = True
        r[1].paragraphs[0].add_run(val)

    if q.get("intent"):
        add_row("Intent", q["intent"])
    if q.get("followups"):
        add_row("Follow-ups", ", ".join(q["followups"][:6]))
    if q.get("good"):
        add_row("What good looks like", q["good"])

    # --- Notes (ruled) ---
    doc.add_paragraph("")
    for _ in range(5):
        _ruled_paragraph(doc)
    doc.add_paragraph("")

def pack_to_docx(
    pack: Dict,
    tenant_name: str = "",
    logo_url: str = "",
    pd_logo_path: str = "assets/powerdash-logo.png",
) -> bytes:
    """
    Expects pack with:
      title, inputs, housekeeping (list[str]), sections (list[{name, notes, questions[]}])
    """
    doc = Document()
    _set_document_defaults(doc)

    # Header
    if logo_url:
        try:
            img = requests.get(logo_url, timeout=6).content
            doc.add_picture(io.BytesIO(img), width=Inches(1.4))
        except Exception:
            pass

    p = doc.add_paragraph()
    r = p.add_run(pack.get("title", "Interview Pack")); r.bold = True; r.font.size = Pt(16)
    meta = f"Interview type: {pack['inputs'].get('interview_type')} · Duration: {pack['inputs'].get('duration_mins')} mins"
    doc.add_paragraph(meta)
    if tenant_name:
        t = doc.add_paragraph(tenant_name); t.runs[0].font.size = Pt(10)

    # Housekeeping
    hk = pack.get("housekeeping") or []
    if hk:
        doc.add_paragraph("")
        h = doc.add_paragraph("Housekeeping"); h.runs[0].bold = True; h.runs[0].font.size = Pt(14)
        for item in hk:
            para = doc.add_paragraph(item)
            try:
                para.style = doc.styles["List Bullet"]
            except Exception:
                pass

    # Sections
    for sec in pack.get("sections", []):
        doc.add_paragraph("")
        s = doc.add_paragraph(sec.get("name", "Section")); s.runs[0].bold = True; s.runs[0].font.size = Pt(14)
        if sec.get("notes"): doc.add_paragraph(sec["notes"])
        for q in (sec.get("questions") or []):
            _add_question_table(doc, q)

    # Footer
    _add_footer_powerdash(doc, pd_logo_path)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

# ==============================
# PDF helpers
# ==============================
def _wrap_lines(c, text: str, width: float, font="Helvetica", size=10):
    words = (text or "").split()
    line, out = "", []
    for w in words:
        t = (line + " " + w).strip()
        if c.stringWidth(t, font, size) > width:
            out.append(line); line = w
        else:
            line = t
    if line: out.append(line)
    return out

def pack_to_pdf(
    pack: Dict,
    tenant_name: str = "",
    logo_url: str = "",
    pd_logo_path: str = "assets/powerdash-logo.png",
) -> bytes:
    """
    Polished PDF with full-width question line, label/value rows below,
    ruled notes, and PD footer logo on every page.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4
    x = 20 * mm
    y = H - 20 * mm

    # Client logo
    if logo_url:
        try:
            c.drawImage(ImageReader(logo_url), x, y - 15 * mm, width=30 * mm, height=15 * mm,
                        preserveAspectRatio=True, mask="auto")
        except Exception:
            pass

    # Title & meta
    c.setFont("Helvetica-Bold", 14); c.drawString(x, y - 18 * mm, pack.get("title", "Interview Pack"))
    c.setFont("Helvetica", 10)
    meta = f"Interview type: {pack['inputs'].get('interview_type')} · Duration: {pack['inputs'].get('duration_mins')} mins"
    c.drawString(x, y - 23 * mm, meta)
    if tenant_name: c.drawString(x, y - 28 * mm, tenant_name)

    cur_y = y - 35 * mm

    def footer():
        fy = 12 * mm
        try:
            if pd_logo_path and os.path.exists(pd_logo_path):
                img = ImageReader(pd_logo_path)
                img_w = 14 * mm; img_h = 14 * mm
                cx = W / 2
                c.drawImage(img, cx - img_w - 12, fy - 3, width=img_w, height=img_h,
                            preserveAspectRatio=True, mask="auto")
                c.setFont("Helvetica-Oblique", 9)
                c.drawString(cx - 12, fy + 3, "Powered by PowerDash HR")
            else:
                c.setFont("Helvetica-Oblique", 9)
                c.drawCentredString(W / 2, fy + 3, "Powered by PowerDash HR")
        except Exception:
            c.setFont("Helvetica-Oblique", 9)
            c.drawCentredString(W / 2, 12 * mm + 3, "Powered by PowerDash HR")

    def ensure_space(lines_needed: int = 12):
        nonlocal cur_y
        if cur_y < (40 * mm + lines_needed * 12):
            footer(); c.showPage(); cur_y = H - 20 * mm

    # ---- QUESTION BLOCK (PDF) ----
    def question_box(q: Dict):
        nonlocal cur_y
        left, right = x, W - x
        top = cur_y

        # QUESTION: full width, bold
        c.setFont("Helvetica-Bold", 11)
        for ln in _wrap_lines(c, (q.get("question") or "").strip(), right - left - 8):
            c.drawString(left + 4, cur_y, ln); cur_y -= 12
        cur_y -= 4

        # Label/value rows
        def row(label, val, label_min=80):
            nonlocal cur_y
            if not val: return
            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 4, cur_y, f"{label}:")
            label_w = max(label_min, c.stringWidth(f"{label}:", "Helvetica-Bold", 10))
            c.setFont("Helvetica", 10)
            for ln in _wrap_lines(c, val, right - left - 8 - label_w):
                c.drawString(left + 4 + label_w, cur_y, ln); cur_y -= 12
            cur_y -= 2

        row("Intent", q.get("intent", ""), label_min=60)
        if q.get("good"): row("What good looks like", q["good"], label_min=140)
        if q.get("followups"): row("Follow-ups", ", ".join(q.get("followups", [])[:6]), label_min=100)

        # Notes (5 dashed lines)
        for _ in range(5):
            c.setDash(1, 3); c.line(left + 4, cur_y, right - 4, cur_y); c.setDash(); cur_y -= 10

        # Draw outer box AFTER content
        box_height = top - cur_y
        c.rect(left, cur_y, right - left, box_height, stroke=1, fill=0)
        cur_y -= 8

    # Housekeeping
    hk = pack.get("housekeeping") or []
    if hk:
        c.setFont("Helvetica-Bold", 12); c.drawString(x, cur_y, "Housekeeping"); cur_y -= 14
        c.setFont("Helvetica", 10)
        for item in hk:
            for ln in _wrap_lines(c, "• " + item, W - 2 * x):
                c.drawString(x, cur_y, ln); cur_y -= 12
        cur_y -= 6
        ensure_space()

    # Sections
    for sec in pack.get("sections", []):
        c.setFont("Helvetica-Bold", 12); c.drawString(x, cur_y, sec.get("name", "Section")); cur_y -= 14
        if sec.get("notes"):
            c.setFont("Helvetica", 10)
            for ln in _wrap_lines(c, sec["notes"], W - 2 * x):
                c.drawString(x, cur_y, ln); cur_y -= 12
            cur_y -= 2

        for q in (sec.get("questions") or []):
            ensure_space(); question_box(q)

        ensure_space()

    footer(); c.save(); buf.seek(0)
    return buf.getvalue()
