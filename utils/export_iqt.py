# utils/export_iqt.py
import io, os, requests
from typing import Dict, List

# ---------- DOCX ----------
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE

# ---------- PDF ----------
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader

FONT_NAME = "Source Sans 3"      # fallback to Calibri if not present
LABEL_COL_IN = 1.2               # DOCX label column width
VALUE_COL_IN = 5.8               # DOCX value column width
NOTES_HEIGHT_PT = 100            # ~6–7 lines of whitespace

# ==============================
# DOCX helpers
# ==============================
def _set_document_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(11)

def _add_footer_powerdash(doc: Document, pd_logo_path: str):
    """Footer on every section → repeats on every page."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        try:
            if pd_logo_path and os.path.exists(pd_logo_path):
                run.add_picture(pd_logo_path, width=Inches(0.22))
                p.add_run("  Powered by PowerDash HR").italic = True
            else:
                p.add_run("Powered by PowerDash HR").italic = True
        except Exception:
            p.add_run("Powered by PowerDash HR").italic = True

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _set_tbl_borders(table, size="8", color="222222"):
    """
    Apply box borders to a python-docx table by manipulating the XML.
    Works across python-docx versions (no .tblBorders attribute access).
    """
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.append(tblPr)

    # remove any existing borders node
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))     # size in eighths of a point (string)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)

    tblPr.append(borders)

def _set_tbl_cell_margins(table, top=160, start=160, bottom=140, end=160):
    """
    Set table cell padding (margins) in twips. Creates <w:tblCellMar> if needed.
    """
    def _marg(tag, val):
        el = OxmlElement(tag)
        el.set(qn("w:w"), str(val))   # twips
        el.set(qn("w:type"), "dxa")
        return el

    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.append(tblPr)

    mar = tblPr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tblPr.append(mar)
    else:
        # clear existing margins to avoid duplicates
        for child in list(mar):
            mar.remove(child)

    mar.append(_marg("w:top", top))
    mar.append(_marg("w:start", start))
    mar.append(_marg("w:bottom", bottom))
    mar.append(_marg("w:end", end))

def _set_row_height(row, points: int, rule=WD_ROW_HEIGHT_RULE.EXACTLY):
    row.height = Pt(points)
    row.height_rule = rule

def _para(p, text="", bold=False, size=11, space_after=4):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def _add_question_table(doc: Document, q: Dict):
    """
    Word layout to mirror PDF:
      - Table with borders (the 'box')
      - Row 1: Question (merged full width), bold, with extra top padding
      - Row 2..n: label/value rows
      - Last row: blank 'notes' cell with fixed height (whitespace, no dots)
    """
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths
    try:
        tbl.columns[0].width = Inches(LABEL_COL_IN)
        tbl.columns[1].width = Inches(VALUE_COL_IN)
    except Exception:
        pass

    _set_tbl_borders(tbl, size="8", color="222222")
    _set_tbl_cell_margins(tbl, top=160, start=160, bottom=140, end=160)  # generous padding

    # --- Question row (merged full width, with extra top space) ---
    first = tbl.rows[0].cells
    try:
        qcell = first[0].merge(first[1])
    except Exception:
        qcell = first[0]
    qp = qcell.paragraphs[0]
    _para(qp, (q.get("question") or "").strip(), bold=True, size=12, space_after=6)
    # Add a tiny blank paragraph to simulate extra top padding visually
    qcell.add_paragraph("")

    # helper for label/value rows
    def add_row(label: str, value: str):
        if not value:
            return
        row = tbl.add_row()
        # label cell
        lp = row.cells[0].paragraphs[0]
        _para(lp, label + ":", bold=True, size=11, space_after=2)
        # value cell
        vp = row.cells[1].paragraphs[0]
        _para(vp, value, bold=False, size=11, space_after=2)

    if q.get("intent"):
        add_row("Intent", q["intent"])
    if q.get("good"):
        add_row("What good looks like", q["good"])
    if q.get("followups"):
        add_row("Follow-ups", ", ".join(q["followups"][:6]))

    # --- Notes whitespace row ---
    notes_row = tbl.add_row()
    # Merge both cells so the space spans full width (cleaner look)
    try:
        notes_cell = notes_row.cells[0].merge(notes_row.cells[1])
    except Exception:
        notes_cell = notes_row.cells[0]
    # Force row height for whitespace area
    _set_row_height(notes_row, NOTES_HEIGHT_PT, rule=WD_ROW_HEIGHT_RULE.EXACTLY)
    notes_cell.paragraphs[0].add_run("")  # keep empty

    # Spacer after the table
    doc.add_paragraph("").paragraph_format.space_after = Pt(8)

def pack_to_docx(
    pack: Dict,
    tenant_name: str = "",
    logo_url: str = "",
    pd_logo_path: str = "assets/powerdash-logo.png",
) -> bytes:
    """
    Expects pack with:
      title, inputs, housekeeping (list[str]),
      sections (list[{name, notes, bullets?, questions[]}])
    """
    doc = Document()
    _set_document_defaults(doc)

    # Header area
    if logo_url:
        try:
            img = requests.get(logo_url, timeout=6).content
            doc.add_picture(io.BytesIO(img), width=Inches(1.4))
        except Exception:
            pass

    title_p = doc.add_paragraph()
    _para(title_p, pack.get("title", "Interview Pack"), bold=True, size=16, space_after=2)
    meta_p = doc.add_paragraph()
    meta = f"Interview type: {pack['inputs'].get('interview_type')} · Duration: {pack['inputs'].get('duration_mins')} mins"
    _para(meta_p, meta, size=11, space_after=0)
    if tenant_name:
        _para(doc.add_paragraph(), tenant_name, size=10, space_after=0)

    # Housekeeping bullets
    hk = pack.get("housekeeping") or []
    if hk:
        doc.add_paragraph("")  # spacer
        _para(doc.add_paragraph(), "Housekeeping", bold=True, size=14, space_after=4)
        for item in hk:
            bp = doc.add_paragraph(item)
            try:
                bp.style = doc.styles["List Bullet"]
            except Exception:
                pass

    # Sections
    for sec in pack.get("sections", []):
        doc.add_paragraph("")  # section spacer
        _para(doc.add_paragraph(), sec.get("name", "Section"), bold=True, size=14, space_after=4)

        # Optional bullets (e.g., Close-down & Next Steps)
        bullets = sec.get("bullets") or []
        if bullets:
            for item in bullets:
                bp = doc.add_paragraph(item)
                try:
                    bp.style = doc.styles["List Bullet"]
                except Exception:
                    pass

        # Optional notes
        if sec.get("notes"):
            _para(doc.add_paragraph(), sec["notes"], size=11, space_after=4)

        # Questions
        for q in (sec.get("questions") or []):
            _add_question_table(doc, q)

    # Footer every page
    _add_footer_powerdash(doc, pd_logo_path)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ==============================
# PDF exporter (already tuned)
# ==============================
def _wrap_lines(c, text: str, width: float, font="Helvetica", size=11):
    words = (text or "").split()
    out, line = [], ""
    for w in words:
        t = (line + " " + w).strip()
        if c.stringWidth(t, font, size) > width:
            if line: out.append(line)
            line = w
        else:
            line = t
    if line: out.append(line)
    return out

def pack_to_pdf(
    pack: Dict,
    tenant_name: str = "",
    logo_url: str = "",
    pd_logo_path: str = "assets/powerdash-logo.png",
) -> bytes:
    """
    Polished PDF with full-width question line, label/value rows below,
    generous WHITE SPACE for notes, and PD footer logo on every page.
    Uses pre-measurement + asymmetric padding to avoid squashing.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4

    # ---- layout constants ----
    MARGIN_X     = 22 * mm
    TOP_Y        = H - 22 * mm
    LINE         = 16          # line height
    PAD_TOP      = 12          # top inset inside box
    PAD_BOTTOM   = 8           # bottom inset
    NOTES_LINES  = 8           # blank lines
    SECTION_GAP  = 10          # gap before section titles
    BLOCK_GAP    = 14          # gap after each box
    BOTTOM_BUF   = 65 * mm     # bottom buffer
    TOP_START_GAP = 12 * mm    # start lower on new page

    x = MARGIN_X
    y = TOP_Y
    cur_y = y

    # ---------- helpers ----------
    def footer():
        fy = 12 * mm
        try:
            if pd_logo_path and os.path.exists(pd_logo_path):
                img = ImageReader(pd_logo_path)
                img_w = 14 * mm; img_h = 14 * mm
                cx = W / 2
                c.drawImage(img, cx - img_w - 12, fy - 3, width=img_w, height=img_h,
                            preserveAspectRatio=True, mask="auto")
                c.setFont("Helvetica-Oblique", 9)
                c.drawString(cx - 12, fy + 3, "Powered by PowerDash HR")
            else:
                c.setFont("Helvetica-Oblique", 9)
                c.drawCentredString(W / 2, fy + 3, "Powered by PowerDash HR")
        except Exception:
            c.setFont("Helvetica-Oblique", 9)
            c.drawCentredString(W / 2, 12 * mm + 3, "Powered by PowerDash HR")

    def _wrap(text, width, font="Helvetica", size=11):
        return _wrap_lines(c, text, width, font=font, size=size)

    def ensure_space(px_needed: float):
        nonlocal cur_y
        if cur_y - px_needed < BOTTOM_BUF:
            footer()
            c.showPage()
            cur_y = TOP_Y - TOP_START_GAP

    # ---------- header ----------
    if logo_url:
        try:
            c.drawImage(ImageReader(logo_url), x, y - 15 * mm, width=30 * mm, height=15 * mm,
                        preserveAspectRatio=True, mask="auto")
        except Exception:
            pass

    c.setFont("Helvetica-Bold", 15); c.drawString(x, y - 18 * mm, pack.get("title", "Interview Pack"))
    c.setFont("Helvetica", 11)
    meta = f"Interview type: {pack['inputs'].get('interview_type')} · Duration: {pack['inputs'].get('duration_mins')} mins"
    c.drawString(x, y - 24 * mm, meta)
    if tenant_name: c.drawString(x, y - 30 * mm, tenant_name)

    cur_y = y - 40 * mm

    # ---------- sections of bullets ----------
    def draw_bullets(title: str, items: List[str]):
        nonlocal cur_y
        if not items: return
        cur_y -= SECTION_GAP
        c.setFont("Helvetica-Bold", 13); c.drawString(x, cur_y, title); cur_y -= LINE
        c.setFont("Helvetica", 11)
        for item in items:
            for ln in _wrap("• " + (item or ""), W - 2 * x, size=11):
                c.drawString(x, cur_y, ln); cur_y -= LINE
        cur_y -= 4

    draw_bullets("Housekeeping", pack.get("housekeeping") or [])

    # ---------- question block ----------
    def question_block(q: Dict):
        nonlocal cur_y
        left, right = x, W - x
        text_width = right - left - (PAD_TOP + PAD_BOTTOM)

        q_lines  = _wrap((q.get("question") or "").strip(), text_width, font="Helvetica-Bold", size=12)
        intent_lines = _wrap(q.get("intent") or "", text_width - 90, size=11)
        good_lines   = _wrap(q.get("good") or "", text_width - 150, size=11)
        fup_text     = ", ".join((q.get("followups") or [])[:6]) if q.get("followups") else ""
        fup_lines    = _wrap(fup_text, text_width - 110, size=11)

        rows_h = len(q_lines)*LINE + 4
        if intent_lines: rows_h += LINE * (len(intent_lines)+1)
        if good_lines:   rows_h += LINE * (len(good_lines)+1)
        if fup_lines:    rows_h += LINE * (len(fup_lines)+1)
        notes_h = NOTES_LINES * LINE
        block_h = PAD_TOP + rows_h + notes_h + PAD_BOTTOM

        ensure_space(block_h + BLOCK_GAP)

        bottom_y = cur_y - block_h
        c.setLineWidth(1)
        c.roundRect(left, bottom_y, right-left, block_h, 6, stroke=1, fill=0)

        ty = cur_y - PAD_TOP

        c.setFont("Helvetica-Bold", 12)
        for ln in q_lines:
            c.drawString(left + PAD_TOP, ty, ln); ty -= LINE
        ty -= 2

        def row(lbl, lines, label_min):
            nonlocal ty
            if not lines: return
            c.setFont("Helvetica-Bold", 11)
            c.drawString(left + PAD_TOP, ty, f"{lbl}:")
            lbl_w = max(label_min, c.stringWidth(f"{lbl}:", "Helvetica-Bold", 11))
            c.setFont("Helvetica", 11)
            for ln in lines:
                c.drawString(left + PAD_TOP + lbl_w, ty, ln); ty -= LINE
            ty -= 2

        if intent_lines: row("Intent", intent_lines, label_min=60)
        if good_lines:   row("What good looks like", good_lines, label_min=150)
        if fup_lines:    row("Follow-ups", fup_lines, label_min=110)

        ty -= notes_h     # white space for notes
        cur_y = bottom_y - BLOCK_GAP

    # ---------- draw sections & questions ----------
    for sec in pack.get("sections", []):
        name = sec.get("name", "Section")
        bullets = sec.get("bullets") or []

        cur_y -= SECTION_GAP
        c.setFont("Helvetica-Bold", 13); c.drawString(x, cur_y, name); cur_y -= LINE

        if bullets:
            c.setFont("Helvetica", 11)
            for item in bullets:
                for ln in _wrap("• " + (item or ""), W - 2*x, size=11):
                    c.drawString(x, cur_y, ln); cur_y -= LINE
            cur_y -= 4

        if sec.get("notes"):
            c.setFont("Helvetica", 11)
            for ln in _wrap(sec["notes"], W - 2*x, size=11):
                c.drawString(x, cur_y, ln); cur_y -= LINE
            cur_y -= 4

        for q in (sec.get("questions") or []):
            question_block(q)

    footer()
    c.save()
    buf.seek(0)
    return buf.getvalue()
